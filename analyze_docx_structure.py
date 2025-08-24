#!/usr/bin/env python3

from docx import Document
import os

def analyze_docx_structure(docx_path):
    """Analyze the actual structure of the Word document to understand layout."""
    
    doc = Document(docx_path)
    
    print("=== WORD DOCUMENT STRUCTURE ANALYSIS ===")
    print(f"Document: {docx_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print()
    
    # Analyze each paragraph with its position
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"Paragraph {i+1}: {para.text[:100]}{'...' if len(para.text) > 100 else ''}")
            
            # Check if this paragraph has any special formatting
            if para.style:
                print(f"  Style: {para.style.name}")
            
            # Check for any runs with special formatting
            for run in para.runs:
                if run.bold or run.italic:
                    print(f"  Run formatting: bold={run.bold}, italic={run.italic}")
            
            print()
    
    # Analyze images
    print("=== IMAGE ANALYSIS ===")
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            print(f"Image {image_count}: {rel.target_ref}")
            print(f"  Target: {rel.target_ref}")
            print(f"  Type: {rel.reltype}")
            print()
    
    print(f"Total images found: {image_count}")

if __name__ == "__main__":
    analyze_docx_structure('docs/word_docs/MODFLOW-GridApproach.docx')
